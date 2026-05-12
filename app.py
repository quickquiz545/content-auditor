from __future__ import annotations

import cgi
import html
import io
import json
import re
import sys
import tempfile
from dataclasses import dataclass, asdict
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Iterable


APP_TITLE = "Content Auditor"
GUIDELINE_SOURCE = "Guidelines for Authors.docx"


GUIDELINE_SUMMARY = [
    "Use AP-style discipline, factual scientific prose, and active voice.",
    "Avoid opinions, analogies, everyday language, weak certainty, and contextless words.",
    "Do not start sentences with if or because, and do not use the word also.",
    "Use one focus, keep the central entity in every section, and answer search intent in the intro.",
    "Write short, complete sentences that follow subject-predicate-object declarations.",
    "Answer question headings directly, with certainty, and without distancing the answer.",
    "Use facts, consensus, research details, examples, data points, units, and percentages.",
    "Introduce abbreviations at first mention, such as Bitcoin (BTC).",
    "Keep featured snippets within 40 words and 320 characters.",
    "Avoid first-paragraph anchors and anchor text as the first words of a paragraph.",
]


BANNED_PHRASES = {
    "Opinion or weak certainty": [
        "i think",
        "in my opinion",
        "pretty sure",
        "i am sure",
        "i'm sure",
        "it might",
        "it may",
        "perhaps",
        "probably",
        "possibly",
    ],
    "Back-reference to earlier content": [
        "as stated before",
        "as mentioned above",
        "as discussed earlier",
        "as explained in",
        "as it is explained",
        "as noted earlier",
    ],
    "Everyday or filler language": [
        "basically",
        "kind of",
        "sort of",
        "a lot of",
        "you know",
        "stuff",
        "things",
        "really",
        "very",
    ],
    "Analogy markers": [
        "think of it as",
        "imagine",
        "just like",
        "similar to a",
        "as if",
    ],
    "Stressful or hateful wording": [
        "you will be dead",
        "hate",
        "hateful",
        "awful",
        "terrible",
        "disaster",
        "catastrophe",
    ],
}

MODAL_WORDS = {
    "can",
    "could",
    "may",
    "might",
    "should",
    "would",
}

TRANSITION_WORDS = {
    "however",
    "therefore",
    "because",
    "although",
    "while",
    "whereas",
    "in contrast",
    "for example",
    "for instance",
}

RESEARCH_WORDS = {
    "study",
    "research",
    "researcher",
    "journal",
    "publication",
    "survey",
    "analysis",
    "data",
    "dataset",
    "consensus",
    "according to",
    "reported",
    "observed",
}

UNIT_PATTERN = re.compile(
    r"\b(?:kg|g|mg|lb|lbs|oz|km|m|cm|mm|mi|ft|in|mph|km/h|kwh|kw|w|v|hz|gb|mb|ms|s|sec|mins?|hours?|days?|years?|usd|rs|lkr|%|percent)\b",
    re.I,
)
YEAR_PATTERN = re.compile(r"\b(?:19|20)\d{2}\b")
NUMBER_PATTERN = re.compile(r"\b\d+(?:[.,]\d+)?\b")
SENTENCE_SPLIT_PATTERN = re.compile(r"(?<=[.!?])\s+")
WORD_PATTERN = re.compile(r"\b[A-Za-z][A-Za-z'-]*\b")
PASSIVE_PATTERN = re.compile(
    r"\b(?:is|are|was|were|be|been|being|has been|have been|had been)\s+\w+(?:ed|en)\b",
    re.I,
)
MARKDOWN_LINK_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\([^)]+\)|\[(?:image|img|figure)[^\]]*\]", re.I)
TABLE_PATTERN = re.compile(r"^\s*\|.+\|\s*$", re.M)


@dataclass
class Issue:
    severity: str
    category: str
    rule: str
    message: str
    suggestion: str
    evidence: str
    line: int | None = None


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - depends on runtime packages
        raise ValueError("DOCX support needs the python-docx package.") from exc

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        doc = Document(str(tmp_path))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return normalize_text("\n\n".join(parts))
    finally:
        try:
            tmp_path.unlink()
        except OSError:
            pass


def extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - depends on runtime packages
        raise ValueError("PDF support needs the pypdf package.") from exc

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return normalize_text("\n\n".join(parts))


def extract_file(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return extract_docx(data)
    if suffix == ".pdf":
        return extract_pdf(data)
    if suffix in {".txt", ".md", ".markdown", ".html", ".htm", ".csv"}:
        return normalize_text(data.decode("utf-8", errors="ignore"))
    if suffix == ".rtf":
        decoded = data.decode("utf-8", errors="ignore")
        decoded = re.sub(r"\\[a-z]+\d* ?", " ", decoded)
        decoded = re.sub(r"[{}]", " ", decoded)
        return normalize_text(decoded)
    raise ValueError("Unsupported file type. Upload .txt, .md, .docx, .pdf, .html, .csv, or .rtf.")


def split_sentences(text: str) -> list[str]:
    chunks = SENTENCE_SPLIT_PATTERN.split(re.sub(r"\s+", " ", text.strip()))
    return [chunk.strip() for chunk in chunks if WORD_PATTERN.search(chunk)]


def split_paragraphs(text: str) -> list[str]:
    return [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]


def word_count(text: str) -> int:
    return len(WORD_PATTERN.findall(text))


def line_number_for(text: str, snippet: str) -> int | None:
    if not snippet:
        return None
    index = text.lower().find(snippet.lower()[:80])
    if index < 0:
        return None
    return text[:index].count("\n") + 1


def excerpt(value: str, limit: int = 220) -> str:
    value = re.sub(r"\s+", " ", value).strip()
    if len(value) <= limit:
        return value
    return value[: limit - 1].rstrip() + "..."


def add_issue(
    issues: list[Issue],
    text: str,
    severity: str,
    category: str,
    rule: str,
    message: str,
    suggestion: str,
    evidence: str,
) -> None:
    issues.append(
        Issue(
            severity=severity,
            category=category,
            rule=rule,
            message=message,
            suggestion=suggestion,
            evidence=excerpt(evidence),
            line=line_number_for(text, evidence),
        )
    )


def infer_sections(text: str) -> list[tuple[str, str]]:
    lines = text.splitlines()
    sections: list[tuple[str, list[str]]] = []
    current_title = "Introduction"
    current_body: list[str] = []

    def looks_like_heading(line: str) -> bool:
        stripped = line.strip().strip("#").strip()
        if not stripped or len(stripped) > 90:
            return False
        if stripped.endswith("."):
            return False
        words = stripped.split()
        if len(words) <= 9 and (
            line.lstrip().startswith("#")
            or stripped.endswith("?")
            or stripped.endswith(":")
            or stripped.istitle()
        ):
            return True
        return False

    for raw in lines:
        line = raw.strip()
        if looks_like_heading(line):
            if current_body:
                sections.append((current_title, current_body))
            current_title = line.strip("# ").rstrip(":")
            current_body = []
        else:
            current_body.append(raw)
    if current_body:
        sections.append((current_title, current_body))
    return [(title, normalize_text("\n".join(body))) for title, body in sections if normalize_text("\n".join(body))]


def ngrams(text: str, size: int = 3) -> set[str]:
    words = [w.lower() for w in WORD_PATTERN.findall(text)]
    return {" ".join(words[i : i + size]) for i in range(max(0, len(words) - size + 1))}


def check_banned_language(text: str, issues: list[Issue]) -> None:
    lowered = text.lower()
    for category, phrases in BANNED_PHRASES.items():
        for phrase in phrases:
            pattern = r"\b" + re.escape(phrase) + r"\b"
            match = re.search(pattern, lowered)
            if match:
                evidence = text[match.start() : match.start() + 180]
                add_issue(
                    issues,
                    text,
                    "high" if category in {"Opinion or weak certainty", "Back-reference to earlier content"} else "medium",
                    "Tone and wording",
                    category,
                    f"Guidelines discourage '{phrase}'.",
                    "Replace it with factual, direct, context-specific wording.",
                    evidence,
                )

    for match in re.finditer(r"\balso\b", text, re.I):
        add_issue(
            issues,
            text,
            "high",
            "Tone and wording",
            "Banned word",
            "The guideline says never use the word 'also'.",
            "Remove it or replace it with a more precise transition.",
            text[match.start() : match.start() + 160],
        )
        break


def check_sentence_rules(text: str, sentences: list[str], issues: list[Issue]) -> None:
    long_sentences = [s for s in sentences if word_count(s) > 25]
    for sentence in long_sentences[:8]:
        add_issue(
            issues,
            text,
            "medium",
            "Sentence quality",
            "Short sentences",
            f"This sentence has {word_count(sentence)} words.",
            "Split it so each sentence gives one clear piece of information.",
            sentence,
        )

    for sentence in sentences:
        first_word = WORD_PATTERN.findall(sentence[:30])
        if first_word and first_word[0].lower() in {"if", "because"}:
            add_issue(
                issues,
                text,
                "high",
                "Sentence quality",
                "Opening conditional",
                f"The sentence starts with '{first_word[0]}'.",
                "Move the conditional or cause clause to the second part of the sentence.",
                sentence,
            )
    for sentence in sentences:
        if PASSIVE_PATTERN.search(sentence):
            add_issue(
                issues,
                text,
                "medium",
                "Sentence quality",
                "Active voice",
                "This sentence appears to use passive voice.",
                "Rewrite with a clear subject performing the action.",
                sentence,
            )

    for sentence in sentences:
        words = {w.lower() for w in WORD_PATTERN.findall(sentence)}
        found = sorted(words & MODAL_WORDS)
        if found:
            add_issue(
                issues,
                text,
                "medium",
                "Certainty",
                "Possibility modal",
                f"Modal wording reduces certainty: {', '.join(found)}.",
                "Use definitive wording unless the heading question uses the same modality.",
                sentence,
            )

    for match in re.finditer(r"(^|\n)\s*[^.\n]{3,90}\b(?:are|include|includes|is)\s*:\s*(?:\n|$)", text, re.I):
        add_issue(
            issues,
            text,
            "high",
            "Sentence quality",
            "Half sentence",
            "This looks like a half sentence before a list.",
            "Use a complete setup sentence, such as 'Benefits of X are listed below:'.",
            match.group(0),
        )


def check_structure(
    text: str,
    paragraphs: list[str],
    sections: list[tuple[str, str]],
    issues: list[Issue],
    central_entity: str,
    search_intent: str,
) -> None:
    intro = "\n\n".join(paragraphs[:2])
    outro = paragraphs[-1] if paragraphs else ""

    if central_entity:
        entity_pattern = re.compile(r"\b" + re.escape(central_entity) + r"\b", re.I)
        if not entity_pattern.search(intro):
            add_issue(
                issues,
                text,
                "high",
                "Semantic focus",
                "Central entity in intro",
                "The central entity is missing from the intro.",
                "Mention the central entity early and process the article through that entity.",
                intro or text[:220],
            )
        missing_sections = [
            title
            for title, body in sections
            if title != "Introduction" and not entity_pattern.search(title + " " + body)
        ]
        if missing_sections:
            add_issue(
                issues,
                text,
                "high",
                "Semantic focus",
                "Central entity in sections",
                "Some sections do not mention the central entity.",
                "Use the central entity or a clear synonym in every section.",
                ", ".join(missing_sections[:6]),
            )

    if search_intent:
        intent_terms = [w.lower() for w in WORD_PATTERN.findall(search_intent) if len(w) > 3]
        intro_terms = {w.lower() for w in WORD_PATTERN.findall(intro)}
        missing = [term for term in intent_terms if term not in intro_terms]
        if missing:
            add_issue(
                issues,
                text,
                "medium",
                "Semantic focus",
                "Search intent in intro",
                "The intro does not clearly cover the central search intent.",
                "Answer the core user intent in the intro before expanding details.",
                intro or text[:220],
            )

    if len(paragraphs) >= 3:
        intro_ngrams = ngrams(intro, 3)
        outro_ngrams = ngrams(outro, 3)
        if intro_ngrams and outro_ngrams and not intro_ngrams.intersection(outro_ngrams):
            add_issue(
                issues,
                text,
                "low",
                "Semantic flow",
                "Intro and outro continuity",
                "The intro and outro do not share meaningful 3-word terms.",
                "Use the same important n-grams in the intro and outro when closing the article.",
                outro,
            )

    for title, body in sections:
        if title.endswith("?") and body:
            first_sentence = split_sentences(body)[:1]
            if first_sentence:
                title_terms = {w.lower() for w in WORD_PATTERN.findall(title) if len(w) > 3}
                answer_terms = {w.lower() for w in WORD_PATTERN.findall(first_sentence[0])}
                overlap = title_terms.intersection(answer_terms)
                if len(overlap) < max(1, min(2, len(title_terms))):
                    add_issue(
                        issues,
                        text,
                        "medium",
                        "Heading answers",
                        "Direct answer",
                        "The first sentence after a question heading may not answer the question directly.",
                        "Start with a direct answer that repeats the core subject and predicate.",
                        title + " " + first_sentence[0],
                    )
        if re.search(r"\b(?:types|benefits|examples|steps|methods|factors|ways)\b", title, re.I):
            first_sentence = split_sentences(body)[:1]
            if first_sentence and not re.search(r"\b(?:is|are|refers to|means|describes|include)\b", first_sentence[0], re.I):
                add_issue(
                    issues,
                    text,
                    "medium",
                    "Heading answers",
                    "List heading definition",
                    "A list-based heading should first define the noun in context.",
                    "Add a precise contextual definition before listing items.",
                    title + " " + first_sentence[0],
                )


def check_evidence_and_detail(text: str, sentences: list[str], issues: list[Issue]) -> dict[str, int]:
    lowered = text.lower()
    research_hits = sum(lowered.count(term) for term in RESEARCH_WORDS)
    year_hits = len(YEAR_PATTERN.findall(text))
    numbers = len(NUMBER_PATTERN.findall(text))
    units = len(UNIT_PATTERN.findall(text))
    percentages = len(re.findall(r"\b\d+(?:\.\d+)?\s*(?:%|percent)\b", text, re.I))
    examples = len(re.findall(r"\b(?:for example|for instance|such as|e\.g\.)\b", text, re.I))

    if len(sentences) >= 8 and research_hits == 0:
        add_issue(
            issues,
            text,
            "medium",
            "Evidence",
            "Research support",
            "No research, study, data, or consensus terms were found.",
            "Integrate research source, organization, date, topic, sample count, or consensus details.",
            text[:220],
        )
    if len(sentences) >= 8 and numbers < 2:
        add_issue(
            issues,
            text,
            "low",
            "Evidence",
            "Numeric detail",
            "The article has limited numeric detail.",
            "Add precise values, counts, ranges, dates, units, or percentages where relevant.",
            text[:220],
        )
    if len(sentences) >= 8 and examples < 2:
        add_issue(
            issues,
            text,
            "low",
            "Evidence",
            "Examples",
            "The article has limited example language.",
            "Use multiple examples to create specific contextual connections.",
            text[:220],
        )

    abbreviation_candidates = re.finditer(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b", text)
    for match in abbreviation_candidates:
        phrase = match.group(1)
        after = text[match.end() : match.end() + 12]
        initials = "".join(part[0] for part in phrase.split()).upper()
        if len(initials) >= 2 and not re.match(r"\s*\([A-Z0-9]{2,8}\)", after):
            add_issue(
                issues,
                text,
                "low",
                "Evidence",
                "First abbreviation",
                f"'{phrase}' may need an abbreviation at first mention.",
                f"If the abbreviation is standard, write '{phrase} ({initials})' on first mention.",
                text[match.start() : match.start() + 160],
            )
            break

    return {
        "research_mentions": research_hits,
        "years": year_hits,
        "numbers": numbers,
        "units": units,
        "percentages": percentages,
        "examples": examples,
    }


def check_special_rules(text: str, paragraphs: list[str], issues: list[Issue]) -> None:
    for match in MARKDOWN_LINK_PATTERN.finditer(text):
        link_text = match.group(1)
        before = text[: match.start()]
        paragraph_index = before.count("\n\n")
        current_para_start = before.rfind("\n\n") + 2
        para_prefix = text[current_para_start : match.start()].strip()
        if paragraph_index < 3:
            add_issue(
                issues,
                text,
                "medium",
                "Anchor terms",
                "Early anchor",
                "Anchor text appears in the first 2-3 paragraphs.",
                "Move internal anchors after the intro context is established.",
                link_text,
            )
        if not para_prefix:
            add_issue(
                issues,
                text,
                "medium",
                "Anchor terms",
                "Paragraph-opening anchor",
                "Anchor text appears as the first words of a paragraph.",
                "Place the anchor after a few contextual words.",
                link_text,
            )

    for match in IMAGE_PATTERN.finditer(text):
        before = text[: match.start()].strip()
        previous_sentence = split_sentences(before)[-1:] or [""]
        if word_count(previous_sentence[0]) < 5:
            add_issue(
                issues,
                text,
                "low",
                "Media",
                "Image introduction",
                "An image marker appears without a clear preceding qualifying sentence.",
                "Add a sentence before the image that connects it to the section context.",
                text[match.start() : match.start() + 120],
            )

    if TABLE_PATTERN.search(text) and re.search(r"\b(?:best|#1|number one|top-ranked|beats all)\b", text, re.I):
        add_issue(
            issues,
            text,
            "medium",
            "Tables",
            "Natural comparison",
            "A table area may rank a brand too aggressively.",
            "Keep comparisons natural and highlight attributes without forcing a first-place rank.",
            text[TABLE_PATTERN.search(text).start() : TABLE_PATTERN.search(text).start() + 220],
        )

    for title, body in infer_sections(text):
        if re.search(r"featured snippet|answer box|snippet", title, re.I):
            words = word_count(body)
            chars = len(body)
            if words > 40 or chars > 320:
                add_issue(
                    issues,
                    text,
                    "high",
                    "Featured snippets",
                    "Snippet length",
                    f"Featured snippet text has {words} words and {chars} characters.",
                    "Keep featured snippets within 40 words and 320 characters.",
                    title + " " + body[:180],
                )

    for paragraph in paragraphs:
        transition_count = sum(1 for term in TRANSITION_WORDS if term in paragraph.lower())
        if word_count(paragraph) > 120 and transition_count == 0:
            add_issue(
                issues,
                text,
                "low",
                "Semantic flow",
                "Long paragraph transitions",
                "A long paragraph has no clear transition markers.",
                "Use precise transition wording when connecting declarations, especially contrast or cautions.",
                paragraph,
            )


def score_from_issues(issues: list[Issue]) -> int:
    penalty = 0
    for issue in issues:
        penalty += {"high": 12, "medium": 7, "low": 3}[issue.severity]
    return max(0, min(100, 100 - penalty))


def audit_content(text: str, central_entity: str = "", search_intent: str = "") -> dict:
    text = normalize_text(text)
    if not text:
        raise ValueError("Add pasted text or upload a content file before running the audit.")

    sentences = split_sentences(text)
    paragraphs = split_paragraphs(text)
    sections = infer_sections(text)
    issues: list[Issue] = []

    check_banned_language(text, issues)
    check_sentence_rules(text, sentences, issues)
    check_structure(text, paragraphs, sections, issues, central_entity.strip(), search_intent.strip())
    detail_metrics = check_evidence_and_detail(text, sentences, issues)
    check_special_rules(text, paragraphs, issues)

    severity_order = {"high": 0, "medium": 1, "low": 2}
    issues.sort(key=lambda issue: (severity_order[issue.severity], issue.category, issue.line or 999999))

    high = sum(1 for issue in issues if issue.severity == "high")
    medium = sum(1 for issue in issues if issue.severity == "medium")
    low = sum(1 for issue in issues if issue.severity == "low")
    sentence_lengths = [word_count(sentence) for sentence in sentences]
    avg_sentence = round(sum(sentence_lengths) / len(sentence_lengths), 1) if sentence_lengths else 0

    return {
        "score": score_from_issues(issues),
        "summary": {
            "words": word_count(text),
            "sentences": len(sentences),
            "paragraphs": len(paragraphs),
            "sections": len(sections),
            "average_sentence_words": avg_sentence,
            "high": high,
            "medium": medium,
            "low": low,
        },
        "detail_metrics": detail_metrics,
        "guidelines": GUIDELINE_SUMMARY,
        "issues": [asdict(issue) for issue in issues],
    }


INDEX_HTML = r"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Content Auditor</title>
  <style>
    :root {
      --bg: #f6f7f9;
      --surface: #ffffff;
      --surface-alt: #fbfcfd;
      --ink: #1d232b;
      --muted: #607083;
      --line: #d9e0e7;
      --accent: #0f766e;
      --accent-2: #285f9f;
      --high: #b42318;
      --medium: #a15c07;
      --low: #3b6d2a;
      --soft-high: #fff0ee;
      --soft-medium: #fff7e8;
      --soft-low: #eef8ec;
      --focus: 0 0 0 3px rgba(15, 118, 110, 0.18);
      --shadow: 0 10px 24px rgba(29, 35, 43, 0.08);
    }

    * { box-sizing: border-box; }

    body {
      margin: 0;
      font-family: Arial, Helvetica, sans-serif;
      background: var(--bg);
      color: var(--ink);
      letter-spacing: 0;
    }

    .app {
      min-height: 100vh;
      display: grid;
      grid-template-rows: auto 1fr;
    }

    header {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 16px;
      padding: 14px 22px;
      background: var(--surface);
      border-bottom: 1px solid var(--line);
    }

    .header-actions {
      display: flex;
      align-items: center;
      justify-content: flex-end;
      gap: 10px;
      min-width: 0;
    }

    .brand {
      display: flex;
      align-items: center;
      gap: 12px;
      min-width: 0;
    }

    .mark {
      width: 34px;
      height: 34px;
      border-radius: 8px;
      background: #163a5f;
      color: #fff;
      display: grid;
      place-items: center;
      font-weight: 700;
    }

    h1 {
      font-size: 18px;
      line-height: 1.2;
      margin: 0;
      white-space: nowrap;
    }

    .source {
      color: var(--muted);
      font-size: 13px;
      overflow: hidden;
      text-overflow: ellipsis;
      white-space: nowrap;
    }

    .pill {
      display: inline-flex;
      align-items: center;
      gap: 6px;
      min-height: 28px;
      border: 1px solid var(--line);
      border-radius: 999px;
      padding: 5px 9px;
      color: #364454;
      background: var(--surface-alt);
      font-size: 12px;
      font-weight: 700;
      white-space: nowrap;
    }

    main {
      display: grid;
      grid-template-columns: minmax(340px, 460px) minmax(0, 1fr);
      gap: 18px;
      padding: 18px;
      max-width: 1480px;
      width: 100%;
      margin: 0 auto;
    }

    .panel {
      background: var(--surface);
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: var(--shadow);
      min-width: 0;
    }

    .input-panel {
      padding: 18px;
      align-self: start;
      position: sticky;
      top: 18px;
    }

    .results-panel {
      min-height: calc(100vh - 88px);
      overflow: hidden;
    }

    label {
      display: block;
      font-weight: 700;
      font-size: 13px;
      margin: 0 0 7px;
    }

    input[type="text"], textarea, select {
      width: 100%;
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 10px 11px;
      font: inherit;
      color: var(--ink);
      background: #fff;
    }

    input[type="text"]:focus, textarea:focus, select:focus, button:focus-visible {
      outline: none;
      box-shadow: var(--focus);
      border-color: var(--accent);
    }

    textarea {
      min-height: 260px;
      resize: vertical;
      line-height: 1.45;
    }

    .field {
      margin-bottom: 14px;
    }

    .file-drop {
      border: 1px dashed #aab7c4;
      background: var(--surface-alt);
      border-radius: 8px;
      padding: 14px;
    }

    .file-drop input {
      width: 100%;
    }

    .help {
      margin-top: 8px;
      color: var(--muted);
      font-size: 12px;
      line-height: 1.35;
    }

    .actions {
      display: flex;
      gap: 10px;
      align-items: center;
      margin-top: 16px;
    }

    button {
      border: 0;
      border-radius: 6px;
      padding: 10px 14px;
      font: inherit;
      font-weight: 700;
      cursor: pointer;
      min-height: 40px;
    }

    .primary {
      background: var(--accent);
      color: #fff;
    }

    .secondary {
      background: #e9eef4;
      color: #25313d;
    }

    button:disabled {
      opacity: 0.6;
      cursor: progress;
    }

    .results-head {
      display: grid;
      grid-template-columns: minmax(150px, 210px) 1fr;
      gap: 16px;
      padding: 18px;
      border-bottom: 1px solid var(--line);
      background: var(--surface-alt);
    }

    .score-wrap {
      display: grid;
      justify-items: center;
      align-content: start;
      gap: 8px;
    }

    .score {
      width: 112px;
      height: 112px;
      border-radius: 50%;
      display: grid;
      place-items: center;
      background: conic-gradient(var(--accent) calc(var(--score, 0) * 1%), #e5eaf0 0);
      position: relative;
    }

    .score::after {
      content: "";
      width: 82px;
      height: 82px;
      border-radius: 50%;
      background: var(--surface);
      position: absolute;
    }

    .score span {
      position: relative;
      z-index: 1;
      font-size: 28px;
      font-weight: 800;
    }

    .score-label {
      color: var(--muted);
      font-size: 12px;
      text-align: center;
    }

    .score-note {
      color: var(--muted);
      font-size: 12px;
      line-height: 1.35;
      text-align: center;
      max-width: 170px;
    }

    .metrics {
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(104px, 1fr));
      gap: 8px;
      align-content: start;
    }

    .metric {
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 9px;
      background: #fff;
      min-height: 64px;
    }

    .metric strong {
      display: block;
      font-size: 20px;
      line-height: 1;
      margin-bottom: 6px;
    }

    .metric span {
      color: var(--muted);
      font-size: 12px;
    }

    .metric.is-alert strong { color: var(--high); }
    .metric.is-ok strong { color: var(--accent); }

    .tabs {
      display: flex;
      gap: 4px;
      padding: 12px 18px 0;
      border-bottom: 1px solid var(--line);
      background: var(--surface);
      overflow-x: auto;
    }

    .tab {
      background: transparent;
      color: var(--muted);
      border-bottom: 3px solid transparent;
      border-radius: 0;
      padding: 10px 11px;
      min-height: 38px;
    }

    .tab.active {
      color: var(--ink);
      border-bottom-color: var(--accent-2);
    }

    .result-tools {
      display: grid;
      grid-template-columns: minmax(160px, 1.2fr) repeat(2, minmax(150px, 0.7fr)) auto;
      gap: 10px;
      align-items: end;
      padding: 14px 18px;
      border-bottom: 1px solid var(--line);
      background: var(--surface);
    }

    .tool-field label {
      color: var(--muted);
      font-size: 11px;
      text-transform: uppercase;
      margin-bottom: 5px;
    }

    .issue-count {
      color: var(--muted);
      font-size: 13px;
      text-align: right;
      min-width: 116px;
      padding-bottom: 10px;
    }

    .content {
      padding: 18px;
    }

    .empty {
      color: var(--muted);
      line-height: 1.5;
      max-width: 720px;
    }

    .issue-list {
      display: grid;
      gap: 10px;
    }

    .insight-row {
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(140px, 1fr));
      gap: 10px;
      margin-bottom: 14px;
    }

    .insight {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 11px;
      background: #fff;
      min-height: 68px;
    }

    .insight strong {
      display: block;
      font-size: 18px;
      margin-bottom: 5px;
    }

    .insight span {
      color: var(--muted);
      font-size: 12px;
    }

    .issue {
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: hidden;
      background: #fff;
    }

    .issue-head {
      display: flex;
      align-items: center;
      gap: 10px;
      padding: 12px 13px;
      border-bottom: 1px solid var(--line);
    }

    .badge {
      font-size: 11px;
      line-height: 1;
      font-weight: 800;
      text-transform: uppercase;
      padding: 6px 7px;
      border-radius: 999px;
      white-space: nowrap;
    }

    .badge.high { color: var(--high); background: var(--soft-high); }
    .badge.medium { color: var(--medium); background: var(--soft-medium); }
    .badge.low { color: var(--low); background: var(--soft-low); }

    .issue-title {
      min-width: 0;
      font-weight: 800;
      overflow-wrap: anywhere;
    }

    .line {
      margin-left: auto;
      color: var(--muted);
      font-size: 12px;
      white-space: nowrap;
    }

    .issue-body {
      display: grid;
      grid-template-columns: minmax(0, 1fr) minmax(220px, 0.8fr);
      gap: 12px;
      padding: 13px;
      line-height: 1.45;
    }

    .message {
      margin: 0 0 9px;
    }

    .suggestion {
      margin: 0;
      color: #2d4f42;
    }

    .evidence {
      background: #f5f7f9;
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 10px;
      color: #354252;
      font-size: 13px;
      overflow-wrap: anywhere;
    }

    .guidelines {
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 10px;
      margin-top: 12px;
    }

    .guideline {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 12px;
      background: #fff;
      line-height: 1.4;
    }

    .empty-state {
      border: 1px dashed var(--line);
      border-radius: 8px;
      padding: 18px;
      background: var(--surface-alt);
    }

    .status {
      color: var(--muted);
      font-size: 13px;
      min-height: 18px;
    }

    .error {
      color: var(--high);
      font-weight: 700;
    }

    @media (max-width: 960px) {
      main {
        grid-template-columns: 1fr;
      }

      .input-panel {
        position: static;
      }

      .results-head {
        grid-template-columns: 1fr;
      }

      .metrics, .guidelines {
        grid-template-columns: repeat(2, minmax(0, 1fr));
      }

      .result-tools {
        grid-template-columns: 1fr 1fr;
      }

      .issue-count {
        grid-column: 1 / -1;
        text-align: left;
        padding-bottom: 0;
      }

      .issue-body {
        grid-template-columns: 1fr;
      }
    }

    @media (max-width: 560px) {
      header {
        align-items: flex-start;
        flex-direction: column;
      }

      .header-actions {
        align-items: flex-start;
        flex-direction: column;
      }

      main {
        padding: 10px;
      }

      .metrics, .guidelines, .result-tools {
        grid-template-columns: 1fr;
      }
    }
  </style>
</head>
<body>
  <div class="app">
    <header>
      <div class="brand">
        <div class="mark">CA</div>
        <div>
          <h1>Content Auditor</h1>
          <div class="source">Rules loaded from Guidelines for Authors.docx</div>
        </div>
      </div>
      <div class="header-actions">
        <span class="pill">Rule pack v1</span>
        <div class="source">Paste content, upload a file, then run an evidence-based audit.</div>
      </div>
    </header>

    <main>
      <section class="panel input-panel">
        <form id="auditForm">
          <div class="field">
            <label for="centralEntity">Central entity</label>
            <input id="centralEntity" name="centralEntity" type="text" placeholder="Example: Laser cutting" />
            <div class="help">Used to check whether each section stays focused on the main entity.</div>
          </div>

          <div class="field">
            <label for="searchIntent">Central search intent</label>
            <input id="searchIntent" name="searchIntent" type="text" placeholder="Example: Does laser cutting produce clean edges?" />
            <div class="help">Used to check whether the intro answers the core user intent.</div>
          </div>

          <div class="field">
            <label for="contentFile">Upload content file</label>
            <div class="file-drop">
              <input id="contentFile" name="contentFile" type="file" accept=".txt,.md,.markdown,.docx,.pdf,.html,.htm,.csv,.rtf" />
              <div class="help">Supported: TXT, Markdown, DOCX, PDF, HTML, CSV, RTF.</div>
            </div>
          </div>

          <div class="field">
            <label for="contentText">Or paste content</label>
            <textarea id="contentText" name="contentText" placeholder="Paste the article or section draft here..."></textarea>
          </div>

          <div class="actions">
            <button class="primary" id="runAudit" type="submit">Run audit</button>
            <button class="secondary" id="clearForm" type="button">Clear</button>
          </div>
          <p class="status" id="status"></p>
        </form>
      </section>

      <section class="panel results-panel">
        <div class="results-head">
          <div class="score-wrap">
            <div class="score" id="scoreDial" style="--score: 0"><span id="scoreValue">--</span></div>
            <div class="score-label">Guideline fit score</div>
            <div class="score-note" id="scoreNote">Waiting for content.</div>
          </div>
          <div class="metrics" id="metrics"></div>
        </div>

        <div class="tabs" id="tabs" aria-label="Issue severity filters">
          <button class="tab active" type="button" data-filter="all">All issues</button>
          <button class="tab" type="button" data-filter="high">High</button>
          <button class="tab" type="button" data-filter="medium">Medium</button>
          <button class="tab" type="button" data-filter="low">Low</button>
          <button class="tab" type="button" data-filter="guidelines">Guidelines</button>
        </div>

        <div class="result-tools" id="resultTools">
          <div class="tool-field">
            <label for="issueSearch">Search findings</label>
            <input id="issueSearch" type="text" placeholder="Rule, category, suggestion, evidence..." />
          </div>
          <div class="tool-field">
            <label for="categoryFilter">Category</label>
            <select id="categoryFilter">
              <option value="all">All categories</option>
            </select>
          </div>
          <div class="tool-field">
            <label for="sortIssues">Sort by</label>
            <select id="sortIssues">
              <option value="severity">Severity</option>
              <option value="category">Category</option>
              <option value="line">Line</option>
            </select>
          </div>
          <div class="issue-count" id="issueCount">No audit yet</div>
        </div>

        <div class="content" id="results">
          <div class="empty-state">
            <p class="empty">Run an audit to see prioritized guideline issues, evidence snippets, and rewrite suggestions.</p>
          </div>
        </div>
      </section>
    </main>
  </div>

  <script>
    const form = document.querySelector("#auditForm");
    const statusEl = document.querySelector("#status");
    const resultsEl = document.querySelector("#results");
    const metricsEl = document.querySelector("#metrics");
    const scoreDial = document.querySelector("#scoreDial");
    const scoreValue = document.querySelector("#scoreValue");
    const scoreNote = document.querySelector("#scoreNote");
    const tabs = document.querySelector("#tabs");
    const runButton = document.querySelector("#runAudit");
    const issueSearch = document.querySelector("#issueSearch");
    const categoryFilter = document.querySelector("#categoryFilter");
    const sortIssues = document.querySelector("#sortIssues");
    const issueCount = document.querySelector("#issueCount");
    let lastResult = null;
    const state = {
      severity: "all",
      query: "",
      category: "all",
      sort: "severity",
    };

    const escapeHtml = (value) => String(value ?? "")
      .replaceAll("&", "&amp;")
      .replaceAll("<", "&lt;")
      .replaceAll(">", "&gt;")
      .replaceAll('"', "&quot;")
      .replaceAll("'", "&#039;");

    function metric(label, value, className = "") {
      return `<div class="metric ${className}"><strong>${escapeHtml(value)}</strong><span>${escapeHtml(label)}</span></div>`;
    }

    function renderMetrics(result) {
      const s = result.summary;
      const d = result.detail_metrics;
      metricsEl.innerHTML = [
        metric("Words", s.words),
        metric("Sentences", s.sentences),
        metric("Avg sentence", s.average_sentence_words),
        metric("High", s.high, s.high ? "is-alert" : "is-ok"),
        metric("Medium", s.medium),
        metric("Low", s.low),
        metric("Research", d.research_mentions),
        metric("Numbers", d.numbers),
        metric("Units", d.units),
      ].join("");
    }

    function scoreMessage(score) {
      if (score >= 85) return "Strong fit. Review remaining details before publishing.";
      if (score >= 65) return "Mostly aligned. Fix medium and high findings next.";
      if (score >= 40) return "Needs revision. Start with high-severity findings.";
      return "Major rewrite needed. Resolve high-severity findings first.";
    }

    function resetFilters() {
      state.severity = "all";
      state.query = "";
      state.category = "all";
      state.sort = "severity";
      issueSearch.value = "";
      categoryFilter.innerHTML = `<option value="all">All categories</option>`;
      sortIssues.value = "severity";
      for (const tab of tabs.querySelectorAll(".tab")) {
        tab.classList.toggle("active", tab.dataset.filter === "all");
      }
      issueCount.textContent = "No audit yet";
    }

    function renderCategoryOptions(result) {
      const categories = [...new Set(result.issues.map((issue) => issue.category))].sort();
      categoryFilter.innerHTML = [
        `<option value="all">All categories</option>`,
        ...categories.map((category) => `<option value="${escapeHtml(category)}">${escapeHtml(category)}</option>`),
      ].join("");
      categoryFilter.value = state.category;
    }

    function issueMatchesQuery(issue, query) {
      if (!query) return true;
      const haystack = [
        issue.severity,
        issue.category,
        issue.rule,
        issue.message,
        issue.suggestion,
        issue.evidence,
        issue.line,
      ].join(" ").toLowerCase();
      return haystack.includes(query);
    }

    function getFilteredIssues(result) {
      const severityRank = { high: 0, medium: 1, low: 2 };
      const query = state.query.trim().toLowerCase();
      const issues = result.issues
        .filter((issue) => state.severity === "all" || issue.severity === state.severity)
        .filter((issue) => state.category === "all" || issue.category === state.category)
        .filter((issue) => issueMatchesQuery(issue, query));

      return issues.sort((a, b) => {
        if (state.sort === "category") {
          return a.category.localeCompare(b.category) || severityRank[a.severity] - severityRank[b.severity];
        }
        if (state.sort === "line") {
          return (a.line || 999999) - (b.line || 999999) || severityRank[a.severity] - severityRank[b.severity];
        }
        return severityRank[a.severity] - severityRank[b.severity] || (a.line || 999999) - (b.line || 999999);
      });
    }

    function renderInsightRow(result) {
      const d = result.detail_metrics;
      return `<div class="insight-row">
        <div class="insight"><strong>${escapeHtml(result.summary.sections)}</strong><span>Detected sections</span></div>
        <div class="insight"><strong>${escapeHtml(d.examples)}</strong><span>Example signals</span></div>
        <div class="insight"><strong>${escapeHtml(d.percentages)}</strong><span>Percentages</span></div>
        <div class="insight"><strong>${escapeHtml(d.years)}</strong><span>Year references</span></div>
      </div>`;
    }

    function renderIssues(result) {
      if (state.severity === "guidelines") {
        issueCount.textContent = `${result.guidelines.length} rules shown`;
        resultsEl.innerHTML = `
          <p class="empty">The audit rules are derived from the attached author guidelines.</p>
          <div class="guidelines">
            ${result.guidelines.map((item) => `<div class="guideline">${escapeHtml(item)}</div>`).join("")}
          </div>
        `;
        return;
      }

      const issues = getFilteredIssues(result);
      issueCount.textContent = `${issues.length} of ${result.issues.length} findings`;
      if (!issues.length) {
        resultsEl.innerHTML = `
          ${renderInsightRow(result)}
          <div class="empty-state">
            <p class="empty">No findings match the current filters.</p>
          </div>
        `;
        return;
      }

      resultsEl.innerHTML = `${renderInsightRow(result)}<div class="issue-list">
        ${issues.map((issue) => `
          <article class="issue">
            <div class="issue-head">
              <span class="badge ${escapeHtml(issue.severity)}">${escapeHtml(issue.severity)}</span>
              <div class="issue-title">${escapeHtml(issue.category)}: ${escapeHtml(issue.rule)}</div>
              ${issue.line ? `<div class="line">Line ${escapeHtml(issue.line)}</div>` : ""}
            </div>
            <div class="issue-body">
              <div>
                <p class="message">${escapeHtml(issue.message)}</p>
                <p class="suggestion">${escapeHtml(issue.suggestion)}</p>
              </div>
              <div class="evidence">${escapeHtml(issue.evidence)}</div>
            </div>
          </article>
        `).join("")}
      </div>`;
    }

    function renderResult(result) {
      lastResult = result;
      scoreDial.style.setProperty("--score", result.score);
      scoreValue.textContent = result.score;
      scoreNote.textContent = scoreMessage(result.score);
      renderMetrics(result);
      renderCategoryOptions(result);
      renderIssues(result);
    }

    form.addEventListener("submit", async (event) => {
      event.preventDefault();
      statusEl.textContent = "Running audit...";
      statusEl.className = "status";
      runButton.disabled = true;

      try {
        const response = await fetch("/audit", {
          method: "POST",
          body: new FormData(form),
        });
        const data = await response.json();
        if (!response.ok) {
          throw new Error(data.error || "The audit failed.");
        }
        statusEl.textContent = "Audit complete.";
        resetFilters();
        renderResult(data);
      } catch (error) {
        statusEl.textContent = error.message;
        statusEl.className = "status error";
      } finally {
        runButton.disabled = false;
      }
    });

    document.querySelector("#clearForm").addEventListener("click", () => {
      form.reset();
      statusEl.textContent = "";
      lastResult = null;
      scoreDial.style.setProperty("--score", 0);
      scoreValue.textContent = "--";
      scoreNote.textContent = "Waiting for content.";
      metricsEl.innerHTML = "";
      resetFilters();
      resultsEl.innerHTML = `<div class="empty-state"><p class="empty">Run an audit to see prioritized guideline issues, evidence snippets, and rewrite suggestions.</p></div>`;
    });

    tabs.addEventListener("click", (event) => {
      const button = event.target.closest(".tab");
      if (!button) return;
      state.severity = button.dataset.filter;
      for (const tab of tabs.querySelectorAll(".tab")) {
        tab.classList.toggle("active", tab === button);
      }
      if (lastResult) renderIssues(lastResult);
    });

    issueSearch.addEventListener("input", () => {
      state.query = issueSearch.value;
      if (lastResult) renderIssues(lastResult);
    });

    categoryFilter.addEventListener("change", () => {
      state.category = categoryFilter.value;
      if (lastResult) renderIssues(lastResult);
    });

    sortIssues.addEventListener("change", () => {
      state.sort = sortIssues.value;
      if (lastResult) renderIssues(lastResult);
    });
  </script>
</body>
</html>"""


class AuditorHandler(BaseHTTPRequestHandler):
    server_version = "ContentAuditor/1.0"

    def log_message(self, fmt: str, *args: object) -> None:
        sys.stderr.write("%s - - [%s] %s\n" % (self.address_string(), self.log_date_time_string(), fmt % args))

    def send_json(self, payload: dict, status: int = 200) -> None:
        body = json.dumps(payload, indent=2).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self) -> None:
        if self.path not in {"/", "/index.html"}:
            self.send_error(404)
            return
        body = INDEX_HTML.encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_POST(self) -> None:
        if self.path != "/audit":
            self.send_error(404)
            return
        try:
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )

            pasted_text = form.getfirst("contentText", "") or ""
            central_entity = form.getfirst("centralEntity", "") or ""
            search_intent = form.getfirst("searchIntent", "") or ""
            uploaded_text = ""

            file_item = form["contentFile"] if "contentFile" in form else None
            if file_item is not None and getattr(file_item, "filename", ""):
                uploaded_text = extract_file(file_item.filename, file_item.file.read())

            content = "\n\n".join(part for part in [uploaded_text, pasted_text] if part.strip())
            self.send_json(audit_content(content, central_entity, search_intent))
        except ValueError as exc:
            self.send_json({"error": str(exc)}, status=400)
        except Exception as exc:
            self.send_json({"error": f"Unexpected audit error: {exc}"}, status=500)


def main(argv: Iterable[str] | None = None) -> int:
    args = list(argv or sys.argv[1:])
    port = int(args[0]) if args else 8000
    server = ThreadingHTTPServer(("127.0.0.1", port), AuditorHandler)
    print(f"{APP_TITLE} running at http://127.0.0.1:{port}")
    print(f"Guideline source: {GUIDELINE_SOURCE}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopping server.")
    finally:
        server.server_close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
